from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image, ImageDraw, ImageFont
import os
from docx2pdf import convert

REPORT_DOCX = os.path.join(os.path.dirname(__file__), '..', 'CareerForge_Project_Report.docx')
REPORT_PDF = os.path.join(os.path.dirname(__file__), '..', 'CareerForge_Project_Report.pdf')
ER_IMAGE = os.path.join(os.path.dirname(__file__), '..', 'er_diagram.png')
WORKFLOW_IMAGE = os.path.join(os.path.dirname(__file__), '..', 'workflow.png')

# The report content (shortened to key sections). For full content, you can expand these strings.
report_content = {
    'title': 'CareerForge Career Preparation Platform - Project Report',
    'chapters': [
        {
            'heading': 'Chapter 1: Introduction',
            'sections': [
                ('1.1 Project Overview', 'CareerForge is a full-stack, web-based career preparation platform that helps students and early-career professionals prepare for technical interviews and job placements.'),
                ('1.2 Problem Statement', 'Job seekers face difficulty transitioning from theoretical knowledge to practical, interview-ready skills; resources are fragmented and lack personalization.'),
                ('1.3 Objectives', '- Centralized platform for learning and interview preparation; - Coding practice with graded assessments; - Resume building and company application tools.'),
                ('1.4 Scope of the Project', 'Includes user auth, courses, coding problems, mock interviews, admin console, and resume builder. Excludes scheduling live interviews and payment gateway.'),
                ('1.5 Methodology', 'Agile iterative development with sprints, prototyping, feedback cycles, and CI automation.'),
                ('1.6 Organization of Report', 'This document covers literature survey, system design, implementation, testing, results, and future scope.'),
            ]
        },
        {
            'heading': 'Chapter 2: Literature Survey',
            'sections': [
                ('2.1 Existing System', 'Existing platforms provide piecemeal functionality (coding platforms, MOOCs, job portals).'),
                ('2.2 Limitations', 'Fragmentation, limited company-focused prep, and poor feedback loops.'),
                ('2.3 Proposed System', 'A unified platform integrating learning, practice, assessments, and company application flows.'),
                ('2.4 Comparative Analysis', 'CareerForge compares favorably by combining features of many platforms into one.'),
            ]
        },
        {
            'heading': 'Chapter 3: System Analysis and Design',
            'sections': [
                ('3.1 Requirement Analysis', 'Functional and non-functional requirements focused on auth, problems, judge, and scalability.'),
                ('3.2 Functional Requirements', 'User registration, problem attempts, assessments, admin tools, resume builder.'),
                ('3.3 Non-Functional Requirements', 'Scalability, performance, security, maintainability, and usability.'),
                ('3.4 Use Case Diagram', 'Actors: User, Admin, Recruiter. Use cases: Login, Attempt Problem, Create Assessment, Manage Content.'),
                ('3.5 Data Flow Diagram', 'High-level flows: User -> Auth -> APIs; Submission -> Judge -> Results -> Dashboard.'),
                ('3.6 ER Diagram', 'Entities: User, Problem, Submission, Assessment, Company, Job. Diagram included as image.'),
                ('3.7 System Architecture', 'React frontend, Spring Boot backend, judge workers, MySQL, optional Redis cache.'),
            ]
        },
        {
            'heading': 'Chapter 4: Implementation',
            'sections': [
                ('4.1 Technologies Used', 'HTML, CSS, JavaScript, React, Monaco, Java Spring Boot, MySQL, Maven, Node.js/npm.'),
                ('4.2 Module Description', 'Auth, User/Profile, Course, Problem, Judge, Assessment, Admin Console, Company Module.'),
                ('4.3 Database Design', 'Relational design with tables for users, problems, testcases, submissions, assessments, companies.'),
                ('4.4 Screenshots of Implementation', 'Screenshots to be included; placeholders present.'),
                ('4.5 Code Snippets', 'Reference to key modules and locations in the repo (backend/src, frontend/src).'),
            ]
        },
        {
            'heading': 'Chapter 5: Testing',
            'sections': [
                ('5.1 Test Plan', 'Unit, integration, frontend unit tests, E2E tests, security testing.'),
                ('5.2 Test Cases', 'Sample test cases for signup, login, attempt problem, submit assessment.'),
                ('5.3 Test Results', 'Run tests and summarize pass/fail counts (to be populated).'),
                ('5.4 Bug Fixes and Validation', 'List and fix issues found during testing; examples included.'),
            ]
        },
        {
            'heading': 'Chapter 6: Results and Discussion',
            'sections': [
                ('6.1 Output Screenshots', 'Screenshots and sample outputs (placeholders).'),
                ('6.2 Performance Analysis', 'Response times and judge scaling notes.'),
                ('6.3 Discussion of Results', 'User feedback, pilot metrics, improvement areas.'),
            ]
        },
        {
            'heading': 'Chapter 7: Conclusion and Future Enhancement',
            'sections': [
                ('7.1 Conclusion', 'CareerForge provides integrated tools for preparation and recruiting; modular architecture supports future enhancements.'),
                ('7.2 Future Scope', 'Add ML recommendations, video interviews, payment, gamification, multi-tenant support.'),
            ]
        },
    ]
}


def create_docx(path):
    doc = Document()
    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(report_content['title'])
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()

    for ch in report_content['chapters']:
        h = doc.add_heading(ch['heading'], level=1)
        for sec_title, sec_text in ch['sections']:
            p = doc.add_heading(sec_title, level=2)
            para = doc.add_paragraph(sec_text)
            para.style.font.size = Pt(11)
            doc.add_paragraph()

    # Insert images placeholders
    doc.add_page_break()
    doc.add_heading('ER Diagram', level=1)
    if os.path.exists(ER_IMAGE):
        doc.add_picture(ER_IMAGE, width=Inches(6))
    else:
        doc.add_paragraph('ER Diagram image is available as er_diagram.png')

    doc.add_page_break()
    doc.add_heading('Workflow Diagram', level=1)
    if os.path.exists(WORKFLOW_IMAGE):
        doc.add_picture(WORKFLOW_IMAGE, width=Inches(6))
    else:
        doc.add_paragraph('Workflow image is available as workflow.png')

    doc.save(path)
    print('Saved DOCX to', path)


def draw_er_diagram(path):
    # Simple ER diagram image using PIL
    img = Image.new('RGBA', (1200, 800), (255, 255, 255, 255))
    d = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype('arial.ttf', 16)
    except Exception:
        font = ImageFont.load_default()

    # Entities - rectangles with labels
    entities = [
        ('User', 50, 50),
        ('Problem', 400, 50),
        ('Submission', 750, 50),
        ('Assessment', 50, 350),
        ('Company', 400, 350),
        ('Job', 750, 350),
    ]

    for name, x, y in entities:
        d.rectangle([x, y, x+250, y+100], outline='black', width=2)
        d.text((x+10, y+10), name, fill='black', font=font)

    # Relationships - lines
    d.line((300,100,400,100), fill='black', width=2)  # User-Problem (example)
    d.line((650,100,750,100), fill='black', width=2)  # Problem-Submission
    d.line((175,150,175,350), fill='black', width=2)  # User-Assessment
    d.line((525,150,525,350), fill='black', width=2)  # Problem-Company
    d.line((900,150,900,350), fill='black', width=2)  # Submission-Job

    img.save(path)
    print('Saved ER diagram to', path)


def draw_workflow(path):
    img = Image.new('RGBA', (1200, 600), (255,255,255,255))
    d = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype('arial.ttf', 16)
    except Exception:
        font = ImageFont.load_default()

    steps = [
        ('User Signup/Login', 50, 50),
        ('Browse Roadmaps', 50, 150),
        ('Attempt Problems', 50, 250),
        ('Take Assessments', 400, 250),
        ('Apply to Companies', 750, 250),
        ('Admin Console', 400, 50),
    ]

    for text, x, y in steps:
        d.rectangle([x, y, x+280, y+60], outline='black', width=2)
        d.text((x+10,y+10), text, fill='black', font=font)

    # arrows/lines
    d.line((190,110,190,150), fill='black', width=2)
    d.line((190,210,190,250), fill='black', width=2)
    d.line((330,280,400,280), fill='black', width=2)
    d.line((680,280,750,280), fill='black', width=2)

    img.save(path)
    print('Saved workflow diagram to', path)


if __name__ == '__main__':
    os.makedirs(os.path.dirname(REPORT_DOCX), exist_ok=True)
    draw_er_diagram(ER_IMAGE)
    draw_workflow(WORKFLOW_IMAGE)
    create_docx(REPORT_DOCX)
    # Convert to PDF
    try:
        convert(REPORT_DOCX, REPORT_PDF)
        print('Converted DOCX to PDF at', REPORT_PDF)
    except Exception as e:
        print('Failed to convert to PDF automatically:', e)
        print('You can convert the DOCX to PDF using Word or another converter.')
